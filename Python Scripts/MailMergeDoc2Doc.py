import concurrent.futures
from docx import Document
import copy
import os

# Very important
script_dir = os.path.dirname(os.path.abspath(__file__))
# If you don't change the directory, you can't run the script properly from a .bat file
os.chdir(script_dir)
# Load the source Word document
source_doc = Document("source.docx")
table = source_doc.tables[0]

# Get the header names from the first row of the table
header_names = [cell.text for cell in table.rows[0].cells]

# Load the docx template file
template = Document("template.docx")

def process_row(row, header_names, template):
    data = {}
    for i, cell in enumerate(row.cells):
        data[header_names[i]] = cell.text

    # Create a new document based on the template
    doc = copy.deepcopy(template)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            new_run_text = run.text
            for header_name in header_names:
                placeholder = "<" + header_name + ">"
                if placeholder in new_run_text:
                    new_run_text = new_run_text.replace(placeholder, str(data.get(header_name, "")))
            run.text = new_run_text

   # Save the output file
    output_folder = "Output"
    if not os.path.exists(output_folder):
    	os.makedirs(output_folder)
    x = f"{output_folder}/{row.cells[0].text}.docx"
    doc.save(x)
    return x




def log_result(future):
    result = future.result()
    print(result)

with concurrent.futures.ThreadPoolExecutor() as executor:
    # Iterate over the rows in the table, starting from the second row
    for row in table.rows[1:]:
        future = executor.submit(process_row, row, header_names, template)
        future.add_done_callback(log_result)
